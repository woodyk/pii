#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: textextract.py
# Author: Wadih Khairallah
# Created: 2024-12-01 12:12:08
# Modified: 2025-05-05

import os
import re
import json
import math
import socket
import platform
import subprocess
import hashlib
import string

from datetime import datetime
from io import StringIO
from urllib.parse import urlparse
from typing import (
    Optional,
    Dict,
    Any,
    List,
)

import magic
import pytesseract
import requests
import pandas as pd
import speech_recognition as sr
import fitz

from bs4 import BeautifulSoup
from collections import Counter
from docx import Document
from mss import mss
from PIL import Image
from pydub import AudioSegment
from rich.console import Console

console = Console()
print = console.print
log = console.log


def clean_path(
    path: str
) -> Optional[str]:
    """
    Normalize and validate a filesystem path.

    Args:
        path (str): Input file or directory path.

    Returns:
        Optional[str]: Absolute path if valid; None otherwise.
    """
    p = os.path.expanduser(path)
    p = os.path.abspath(p)
    if os.path.isfile(p) or os.path.isdir(p):
        return p
    return None


def get_screenshot() -> str:
    """
    Capture a full-screen screenshot and save to a temporary file.

    Returns:
        str: File path of saved screenshot PNG.
    """
    tmp_path = "/tmp/sym_screenshot.png"
    with mss() as sct:
        monitor = {"top": 0, "left": 0, "width": 0, "height": 0}
        for mon in sct.monitors:
            monitor["left"] = min(mon["left"], monitor["left"])
            monitor["top"] = min(mon["top"], monitor["top"])
            monitor["width"] = max(
                mon["width"] + mon["left"] - monitor["left"],
                monitor["width"]
            )
            monitor["height"] = max(
                mon["height"] + mon["top"] - monitor["top"],
                monitor["height"]
            )
        screenshot = sct.grab(monitor)
    img = Image.frombytes(
        "RGB", screenshot.size, screenshot.bgra, "raw", "BGRX"
    )
    img_gray = img.convert("L")
    img_gray.save(tmp_path)
    return tmp_path


def extract_exif(
    file_path: str
) -> Optional[Dict[str, Any]]:
    """
    Extract EXIF metadata from a file using exiftool.

    Args:
        file_path (str): Path to the target file.

    Returns:
        Optional[Dict[str, Any]]: Parsed EXIF data, or None on failure.
    """
    exif_data: Optional[Dict[str, Any]] = None
    try:
        result = subprocess.run(
            ['exiftool', '-j', file_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        if result.returncode == 0:
            exif_data = json.loads(result.stdout.decode())[0]
    except Exception as e:
        print(f"Exiftool failed: {e}")
    return exif_data


def text_from_url(
    url: str
) -> Optional[str]:
    """
    Fetch and extract visible text from a web page.

    Args:
        url (str): The target webpage URL.

    Returns:
        Optional[str]: Extracted text, or None on failure.
    """
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(
            ["script", "style", "noscript", "iframe",
             "header", "footer", "meta", "link"]
        ):
            tag.decompose()
        return soup.get_text(separator=" ").strip()
    except requests.RequestException as e:
        print(f"Error fetching URL: {url} - {e}")
        return None


def extract_text(
    file_path: str
) -> Optional[str]:
    """
    Extract text content from a file based on MIME type.

    Supports text, JSON, XML, CSV, Excel, PDF, DOCX, images, audio.

    Args:
        file_path (str): Path to the input file.

    Returns:
        Optional[str]: Extracted text, or None if unsupported or error.
    """
    TEXT_MIME_TYPES = {
        # programming, config, data types...
        "application/json", "application/xml", "application/x-yaml",
        "application/x-toml", "application/x-csv", "application/x-markdown",
        # add others as needed
    }

    path = clean_path(file_path)
    if not path:
        print(f"No such file: {file_path}")
        return None

    mime_type = magic.from_file(path, mime=True)
    try:
        if mime_type.startswith("text/") or mime_type in TEXT_MIME_TYPES:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        elif mime_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            content = text_from_excel(path)
        elif mime_type == "application/pdf":
            content = text_from_pdf(path)
        elif mime_type == \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = text_from_docx(path)
        elif mime_type == "application/msword":
            content = text_from_doc(path)  # legacy .doc
        elif mime_type.startswith("image/"):
            content = text_from_image(path)
        elif mime_type.startswith("audio/"):
            content = text_from_audio(path)
        else:
            content = text_from_other(path)

        if content:
            return content
        else:
            print(f"No content found for file: {path}")
            return None
    except Exception as e:
        print(f"Error reading {path}: {e}")
        return None


def text_from_audio(
    audio_file: str
) -> Optional[str]:
    """
    Transcribe audio to text via Google Speech Recognition.

    Args:
        audio_file (str): Path to the audio file.

    Returns:
        Optional[str]: Transcription, or None on error.
    """
    def audio_to_wav(
        file_path: str
    ) -> str:
        _, ext = os.path.splitext(file_path)
        ext = ext.lstrip('.')
        audio = AudioSegment.from_file(file_path, format=ext)
        wav_path = file_path.replace(f".{ext}", ".wav")
        audio.export(wav_path, format='wav')
        return wav_path

    _, ext = os.path.splitext(audio_file)
    if ext.lower() not in ['.wav', '.wave']:
        audio_file = audio_to_wav(audio_file)
    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(audio_file) as source:
            audio = recognizer.record(source)
        return recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        print("Could not understand audio")
    except sr.RequestError as e:
        print(f"Speech service error: {e}")
    return None


def downloadImage(
    url: str
) -> Optional[str]:
    """
    Download an image from a URL to /tmp/ and return its path.

    Args:
        url (str): Remote image URL.

    Returns:
        Optional[str]: Local file path, or None on failure.
    """
    if is_image(url):
        filename = os.path.basename(urlparse(url).path)
        save_path = os.path.join("/tmp/", filename)
        resp = requests.get(url, stream=True)
        resp.raise_for_status()
        with open(save_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
        return clean_path(save_path)
    print(f"Unable to pull image from {url}")
    return None


def is_image(
    file_path_or_url: str
) -> bool:
    """
    Determine if the given path/URL points to an image.

    Args:
        file_path_or_url (str): Local path or URL.

    Returns:
        bool: True if MIME type starts with 'image/'.
    """
    try:
        mime = magic.from_file(file_path_or_url, mime=True)
        return mime.startswith("image/")
    except Exception:
        return False


def text_from_pdf(
    pdf_path: str
) -> Optional[str]:
    """
    Extract text and image OCR from a PDF using PyMuPDF.

    Args:
        pdf_path (str): Path to PDF file.

    Returns:
        Optional[str]: Combined text and OCR results, or None on error.
    """
    plain_text = ""
    try:
        doc = fitz.open(pdf_path)
        # metadata
        for k, v in doc.metadata.items():
            plain_text += f"{k}: {v}\n"
        for i in range(len(doc)):
            page = doc.load_page(i)
            plain_text += f"\n--- Page {i+1} ---\n"
            txt = page.get_text()
            plain_text += txt or "[No text]\n"
            for img_index, img in enumerate(page.get_images(full=True), start=1):
                xref = img[0]
                base = doc.extract_image(xref)
                img_bytes = base["image"]
                img_path = f"/tmp/page{i+1}-img{img_index}.png"
                with open(img_path, "wb") as img_file:
                    img_file.write(img_bytes)
                ocr = text_from_image(img_path) or ""
                plain_text += f"\n[Image {img_index} OCR]\n{ocr}\n"
        doc.close()
        return plain_text
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return None


def text_from_doc(
    filepath: str,
    min_length: int = 4
) -> str:
    """
    Extract readable strings and metadata from binary Word (.doc) files.

    Args:
        filepath (str): Path to .doc file.
        min_length (int): Minimum string length to extract.

    Returns:
        str: Metadata and text content.
    """
    def extract_printable_strings(
        data: bytes
    ) -> List[str]:
        pattern = re.compile(
            b'[' + re.escape(bytes(string.printable, 'ascii')) +
            b']{%d,}' % min_length
        )
        found = pattern.findall(data)
        return list(dict.fromkeys(m.decode(errors='ignore').strip()
                                   for m in found))

    def clean_strings(
        strs: List[str]
    ) -> List[str]:
        cleaned: List[str] = []
        skip = ["HYPERLINK", "OLE2", "Normal.dotm"]
        for line in strs:
            if any(line.startswith(pref) for pref in skip):
                continue
            cleaned.append(re.sub(r'\s+', ' ', line).strip())
        return cleaned

    with open(filepath, 'rb') as f:
        data = f.read()
    strings = extract_printable_strings(data)
    strings = clean_strings(strings)
    return "\n".join(strings)


def text_from_docx(
    file_path: str
) -> Optional[str]:
    """
    Extract text, tables, and OCR images from a DOCX file.

    Args:
        file_path (str): Path to the .docx file.

    Returns:
        Optional[str]: Combined document text, or None on error.
    """
    path = clean_path(file_path)
    if not path:
        return None
    doc = Document(path)
    plain_text = ""
    try:
        for p in doc.paragraphs:
            if p.text.strip():
                plain_text += p.text.strip() + "\n"
        for tbl in doc.tables:
            plain_text += "\n[Table]\n"
            for row in tbl.rows:
                plain_text += "\t".join(c.text.strip()
                                        for c in row.cells) + "\n"
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                blob = doc.part.rels[rel].target_part.blob
                img_path = f"/tmp/docx_img_{rel}.png"
                with open(img_path, "wb") as img_f:
                    img_f.write(blob)
                ocr = text_from_image(img_path) or ""
                plain_text += f"\n[Image OCR]\n{ocr}\n"
        return plain_text
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return None


def text_from_excel(
    file_path: str
) -> str:
    """
    Convert an Excel workbook to CSV text.

    Args:
        file_path (str): Path to the Excel file.

    Returns:
        str: CSV-formatted string.
    """
    path = clean_path(file_path)
    if not path:
        return ""
    try:
        df = pd.read_excel(path)
        out = StringIO()
        df.to_csv(out, index=False)
        return out.getvalue()
    except Exception as e:
        print(f"Failed Excel -> CSV: {e}")
        return ""


def text_from_image(
    file_path: str
) -> Optional[str]:
    """
    Perform OCR on an image file.

    Args:
        file_path (str): Path to the image.

    Returns:
        Optional[str]: Extracted text, or None on error.
    """
    path = clean_path(file_path)
    if not path:
        return None
    try:
        with Image.open(path) as img:
            txt = pytesseract.image_to_string(img).strip()
            return txt or ""
    except Exception as e:
        print(f"Failed image OCR: {e}")
        return None


def text_from_other(
    file_path: str
) -> Optional[str]:
    """
    Handle unknown file types by reporting stats and metadata.

    Args:
        file_path (str): Path to the file.

    Returns:
        Optional[str]: Plain-text report, or None on error.
    """
    path = clean_path(file_path)
    if not path:
        return None
    try:
        stats = os.stat(path)
        info = {
            "path": path,
            "size": stats.st_size,
            "created": datetime.fromtimestamp(stats.st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(stats.st_mtime).isoformat(),
        }
        return "\n".join(f"{k}: {v}" for k, v in info.items())
    except Exception as e:
        print(f"Error on other file: {e}")
        return None


def extract_metadata(
    file_path: str
) -> Dict[str, Any]:
    """
    Extract comprehensive metadata from any file type.

    Args:
        file_path (str): Path to target file.

    Returns:
        Dict[str, Any]: Nested metadata structure.
    """
    path = clean_path(file_path)
    if not path:
        return {"error": "File not found"}
    meta: Dict[str, Any] = {}
    try:
        stats = os.stat(path)
        meta["size_bytes"] = stats.st_size
        meta["mime"] = magic.from_file(path, mime=True)
        meta["hashes"] = {
            "md5": hashlib.md5(open(path,'rb').read()).hexdigest()}
    except Exception as e:
        meta["error"] = str(e)
    return meta


def main() -> None:
    """
    CLI entry point for text or metadata extraction.
    Parses arguments and prints results.
    """
    import argparse
    parser = argparse.ArgumentParser(
        description="Extract text or metadata from a file"
    )
    parser.add_argument(
        "file",
        type=str,
        help="Path to the input file"
    )
    parser.add_argument(
        "--metadata",
        action="store_true",
        help="Extract metadata instead of text"
    )
    args = parser.parse_args()
    if args.metadata:
        data = extract_metadata(args.file)
        print(json.dumps(data, indent=2))
    else:
        txt = extract_text(args.file)
        print(txt or "No text extracted.")


if __name__ == "__main__":
    main()

