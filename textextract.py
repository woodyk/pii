#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: textextract.py
# Author: Wadih Khairallah
# Description: 
# Created: 2024-12-01 12:12:08
# Modified: 2025-04-28 14:50:03

import json
import math
import re
import os
import subprocess
import string
import magic
import hashlib
import pytesseract
import requests 
import pandas as pd
import speech_recognition as sr
import fitz
import socket
import platform

from bs4 import BeautifulSoup
from collections import Counter
from docx import Document
from datetime import datetime
from mss import mss
from urllib.parse import urlparse
from io import StringIO
from PIL import Image
from pydub import AudioSegment
from rich.console import Console

console = Console()
print = console.print
log = console.log

def clean_path(path):
    path = os.path.expanduser(path)
    path = os.path.abspath(path)
    
    if os.path.isfile(path) or os.path.isdir(path):
        return path
    else:
        return False

def get_screenshot():
    ''' Take screenshot and return text object for all text found in the image '''
    # Screenshot storage path
    path = r'/tmp/sym_screenshot.png'

    with mss() as sct:
        monitor = {"top": 0, "left": 0, "width": 0, "height": 0}
        
        for mon in sct.monitors:
            # get furthest left point
            monitor["left"] = min(mon["left"], monitor["left"])
            # get highest point
            monitor["top"] = min(mon["top"], monitor["top"])
            # get furthest right point
            monitor["width"] = max(mon["width"]+mon["left"]-monitor["left"], monitor["width"])
            # get lowest point
            monitor["height"] = max(mon["height"]+mon["top"]-monitor["top"], monitor["height"])
        
        screenshot = sct.grab(monitor)

    img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
    img_gray = img.convert("L")
    img_gray.save(path)

    return path

def extract_exif(file_path):
    exif_data = None
    try:
        result = subprocess.run(['exiftool', '-j', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0:
            exif_data = json.loads(result.stdout.decode())[0]

    except Exception as e:
        print(f"Exiftool failed: {e}")

    return exif_data

def text_from_url(url):
    """
    Fetch and extract text from a given URL.

    Args:
        url (str): The website URL.

    Returns:
        str: Extracted plain text from the web page.
    """
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Extract only visible text (excluding script, style, and metadata)
        for tag in soup(["script", "style", "noscript", "iframe", "header", "footer", "meta", "link"]):
            tag.decompose()

        text = soup.get_text(separator=" ")
        return text.strip()

    except requests.RequestException as e:
        print(f"Error fetching URL: {url} - {e}")
        return None

def extract_text(file_path):
    """
    Extracts text content from a file based on its MIME type.

    Processes various file types (text, Excel, PDF, Word, images,
    audio) and returns their extracted text as a UTF-8 string. Uses
    helper methods for non-text files (e.g., PDF to markdown,
    audio transcription). Logs issues and returns None if no
    content is found or an error occurs.

    Args:
        file_path (str): Path to the file, cleaned via `clean_path`.

    Returns:
        str or None: Extracted text if successful, else None.

    Raises:
        Exception: Caught internally, logged with `log` function.

    Supported MIME Types:
        - Text: 'text/*', 'application/json', 'application/xml',
          'application/x-yaml', 'text/markdown'
        - Excel: 'application/vnd.ms-excel',
          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        - PDF: 'application/pdf'
        - Word: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        - Image: 'image/*'
        - Audio: 'audio/*'

    Examples:
        >>> extract_text("example.txt")
        'Hello, world!'
        >>> extract_text("example.pdf")
        '# Page 1\\nContent here...'
        >>> extract_text("invalid_file.txt")
        None  # Logs "Error reading invalid_file.txt: [error]"
    """
    TEXT_MIME_TYPES = {
        # ─── Programming Languages ───
        "application/x-python-code",
        "application/x-java-source",
        "application/x-c",
        "application/x-c++",
        "application/x-rust",
        "application/x-go",
        "application/x-haskell",
        "application/x-kotlin",
        "application/x-scala",
        "application/x-lua",
        "application/x-swift",

        # ─── Web and Scripting ───
        "application/javascript",
        "application/x-javascript",
        "application/x-httpd-php",
        "application/x-perl",
        "application/x-ruby",
        "application/x-sh",
        "application/x-shellscript",

        # ─── Config/Markup/Data ───
        "application/json",
        "application/xml",
        "application/x-yaml",
        "application/x-toml",
        "application/x-properties",
        "application/x-ini",
        "application/x-config",
        "application/x-env",

        # ─── SQL and Structured Data ───
        "application/sql",
        "application/x-sql",
        "application/x-csv",
        "application/x-turtle",
        "application/sparql-query",

        # ─── Lightweight Markup ───
        "application/x-latex",
        "application/x-tex",
        "application/x-markdown",
        "application/x-restructuredtext",

        # ─── Certs and Keys ───
        "application/x-pem-file",
        "application/pem-certificate-chain",
        "application/x-pkcs7-certificates",

        # ─── Miscellaneous ───
        "application/x-subrip",
        "application/x-readme",
        "application/x-crontab",
    }


    file_path = clean_path(file_path)
    if not file_path: 
        print(f"No such file: {file_path}")
        return f"No such file: {file_path}"

    file_path = clean_path(file_path)
    mime_type = magic.from_file(file_path, mime=True)
    try:
        content = "" 
        if mime_type.startswith('text/') or mime_type in TEXT_MIME_TYPES:
            with open(file_path, 'r') as f:
                content = f.read()

        elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            content = text_from_excel(file_path)

        elif mime_type == 'application/pdf':
            content = text_from_pdf(file_path)

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content = text_from_docx(file_path)

        elif mime_type == 'application/msword':
            content = text_from_doc(file_path)

        elif mime_type.startswith('image/'):
            content = text_from_image(file_path)

        elif mime_type.startswith('audio/'):
            content = text_from_audio(file_path)
        
        else:
            content = text_from_other(file_path)

        if content is not None and len(content) > 0:
            content = content.encode('utf-8').decode('utf-8', errors='ignore')
            return content

        else:
            print(f"No content found for file: {file_path}")
            return None

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def text_from_audio(audio_file):
    text = ""

    def audio_to_wav(file_path):
        # Extract the file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lstrip('.')

        # Use pydub to convert to WAV
        audio = AudioSegment.from_file(file_path, format=ext)
        wav_file_path = file_path.replace(ext, 'wav')
        audio.export(wav_file_path, format='wav')

        return wav_file_path

    recognizer = sr.Recognizer()
    _, ext = os.path.splitext(audio_file)
    # Convert the file to WAV if necessary
    if ext.lower() not in ['.wav', '.wave']:
        audio_file = audio_to_wav(audio_file)
    try:
        with sr.AudioFile(audio_file) as source:
            audio = recognizer.record(source)

        text = recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        print("Google Speech Recognition could not understand audio")
        return None
    except sr.RequestError as e:
        print(f"Could not request results from Google Speech Recognition service; {e}")
        return None

    return text

def downloadImage(url):
    if is_image(url):
        filename = os.path.basename(urlparse(url).path)
        save_path = os.path.join('/tmp/', filename)

        response = requests.get(url, stream=True)
        response.raise_for_status()

        with open(save_path, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)

        return clean_path(save_path)
    else:
        print(f"Unable to pull image from {url}")
        return None

def is_image(file_path_or_url):
    try:
        mime = magic.from_file(file_path_or_url, mime=True)
        if mime.startswith("image/"):
            return True
    except Exception as e:
        return False

def text_from_pdf(pdf_path):
    """
    Extracts plain text from a PDF using PyMuPDF (fitz),
    including metadata and OCR for images.
    """
    plain_text = ""

    try:
        doc = fitz.open(pdf_path)

        # Extract metadata
        metadata = doc.metadata
        if metadata:
            for key, value in metadata.items():
                plain_text += f"{key}: {value}\n"
            plain_text += "\n"

        # Iterate through pages
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            plain_text += f"\n\n--- Page {page_num + 1} ---\n\n"

            # Extract text
            text = page.get_text()
            plain_text += text if text.strip() else "[No text found on this page]\n"

            # Extract and OCR images
            image_list = page.get_images(full=True)
            if image_list:
                for image_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_path = f"/tmp/page-{page_num + 1}-image-{image_index + 1}.png"

                    # Save image
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)

                    # Perform OCR on the image
                    image_text = text_from_image(image_path)
                    plain_text += f"\n[Extracted Text from Image {image_index + 1}]\n{image_text}\n"
            else:
                plain_text += "\n[No images found]\n"

        doc.close()

    except Exception as e:
        print(f"Error processing PDF with PyMuPDF: {pdf_path}\n{e}")
        return None

    return plain_text

def text_from_doc(filepath, min_length=4):
    def extract_printable_strings(binary_data):
        pattern = re.compile(b'[' + re.escape(bytes(string.printable, 'ascii')) + b']{%d,}' % min_length)
        matches = pattern.findall(binary_data)
        decoded = [m.decode(errors='ignore').strip() for m in matches]
        return list(dict.fromkeys(decoded))  # preserve order, remove duplicates

    def clean_strings(strings):
        cleaned = []
        skip_prefixes = ["HYPERLINK", "OLE2", "bjbj", "Normal.dotm"]
        for line in strings:
            if any(line.startswith(prefix) for prefix in skip_prefixes):
                continue
            line = re.sub(r'HYPERLINK\s+"[^"]+"', '', line)
            line = re.sub(r'\s+', ' ', line).strip()
            if line:
                cleaned.append(line)
        return cleaned

    def extract_metadata(strings):
        metadata = {}
        for line in strings:
            if "Microsoft Word" in line:
                metadata["Software"] = line
            elif "Word.Document" in line:
                metadata["Format"] = line
            elif "Title" in line and "Title" not in metadata:
                metadata["Title"] = line
            elif "MSWordDoc" in line:
                metadata["Format Version"] = line
            elif "Woody K" in line and "Author" not in metadata:
                metadata["Author"] = line
        return metadata

    def remove_non_content_tail(strings):
        cutoff_patterns = ["[Content_Types].xml", "<?xml", "theme/", "PK", "<ds:", "<a:", "_rels/"]
        clean_until = len(strings)
        for i, line in enumerate(strings):
            if any(pat in line for pat in cutoff_patterns):
                clean_until = i
                break
        return strings[:clean_until]

    def group_into_paragraphs(lines):
        paragraphs = []
        paragraph = []
        for line in lines:
            if re.match(r'^[A-Z][a-z]+.*[.:;]$', line) or len(line.split()) > 10:
                paragraph.append(line)
            else:
                if paragraph:
                    paragraphs.append(" ".join(paragraph))
                    paragraph = []
                paragraphs.append(line)
        if paragraph:
            paragraphs.append(" ".join(paragraph))
        return paragraphs

    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File '{filepath}' does not exist.")

    with open(filepath, 'rb') as f:
        binary_data = f.read()

    raw_strings = extract_printable_strings(binary_data)
    cleaned = clean_strings(raw_strings)
    metadata = extract_metadata(cleaned)
    cleaned = remove_non_content_tail(cleaned)
    paragraphs = group_into_paragraphs(cleaned)

    # Create output string
    output = ["--- Metadata ---"]
    for k, v in metadata.items():
        output.append(f"{k}: {v}")
    output.append("\n--- Document Content ---\n")
    output.append("\n\n".join(paragraphs))

    return "\n".join(output)


def text_from_docx(file_path):
    """
    Extracts plain text from a Word (.docx) file, including text, tables, and images with OCR.
    """
    file_path = clean_path(file_path)
    doc = Document(file_path)
    plain_text = ""

    # Extract text from paragraphs
    try:
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                plain_text += paragraph.text.strip() + "\n\n"
    except Exception as e:
        print(f"Error extracting text from Word file: {file_path}\n{e}")
        return None

    # Extract text from tables
    try:
        for table_num, table in enumerate(doc.tables, start=1):
            plain_text += f"\n[Table {table_num}]\n"
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                plain_text += "\t".join(cells) + "\n"
    except Exception as e:
        print(f"Error extracting tables from Word file: {file_path}\n{e}")
        return None

    # Extract and process images
    try:
        image_num = 0
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_num += 1
                image_data = doc.part.rels[rel].target_part.blob  # Extract image data
                image_path = f"/tmp/word_image_{image_num}.png"
                
                # Save image
                with open(image_path, "wb") as img_file:
                    img_file.write(image_data)

                # Perform OCR on the extracted image
                image_text = text_from_image(image_path)
                plain_text += f"\n[Extracted Text from Image {image_num}]\n{image_text}\n"
    except Exception as e:
        print(f"Error extracting images from Word file: {file_path}\n{e}")
        return None

    return plain_text

def text_from_excel(file_path):
    """
    Converts an Excel file to CSV format.
    """
    file_path = clean_path(file_path)
    csv_content = ""
    try:
        df = pd.read_excel(file_path)
        output = StringIO()
        df.to_csv(output, index=False)
        csv_content = output.getvalue()
    except Exception as e:
        print(f"Failed to convert Excel to CSV: {e}")

    return csv_content

def text_from_image(file_path):
    """
    Extracts plain text from an image using OCR.
    """
    file_path = clean_path(file_path)
    try:
        with Image.open(file_path) as img:
            # Perform OCR to extract text
            extracted_text = pytesseract.image_to_string(img).strip()
            return extracted_text if extracted_text else ""

    except Exception as e:
        print(f"Failed to process image: {file_path}, Error: {e}")
        return None

def text_from_other(file_path):
    """
    Extracts information from a file of unknown or unsupported type and returns plain text output.
    """
    file_path = clean_path(file_path)
    file_stats = os.stat(file_path)
    creation_time = datetime.fromtimestamp(file_stats.st_ctime)
    modified_time = datetime.fromtimestamp(file_stats.st_mtime)
    
    file_info = {
        "File Path": str(file_path),
        "File Size (bytes)": str(file_stats.st_size),
        "Creation Time": str(creation_time),
        "Modification Time": str(modified_time),
        "Permissions": oct(file_stats.st_mode & 0o777),
        "MIME Type": str(magic.from_file(file_path, mime=True)),
        "Hashes": {},
        "Readable Strings": [],
        "Magic Numbers": None,
        "Embedded URLs": [],
        "Entropy": None,
        "Exif Data": {},
    }

    # Get EXIF data
    exif_data = extract_exif(file_path)
    if exif_data:
        for key, value in exif_data.items():
            file_info["Exif Data"][key] = value

    def calculate_entropy(data):
        """Calculate Shannon entropy to assess randomness in the file."""
        if not data:
            return "0"
        counter = Counter(data)
        length = len(data)
        entropy = -sum((count / length) * math.log2(count / length) for count in counter.values())
        return str(entropy)

    def extract_strings(data):
        """Extract readable ASCII and Unicode strings."""
        ascii_regex = re.compile(rb'[ -~]{4,}')  # ASCII strings of length >= 4
        unicode_regex = re.compile(rb'(?:[\x20-\x7E][\x00]){4,}')  # Unicode UTF-16 strings
        strings = []
        strings.extend(match.decode('ascii') for match in ascii_regex.findall(data))
        strings.extend(match.decode('utf-16', errors='ignore') for match in unicode_regex.findall(data))
        return strings

    # Read the file as binary
    try:
        with open(file_path, 'rb') as file:
            binary_data = file.read()
            file_info["Hashes"]["SHA-256"] = hashlib.sha256(binary_data).hexdigest()
            file_info["Hashes"]["MD5"] = hashlib.md5(binary_data).hexdigest()
            file_info["Readable Strings"] = extract_strings(binary_data)[:10]  # Limit to 10 strings
            file_info["Entropy"] = calculate_entropy(binary_data)
            file_info["Magic Numbers"] = binary_data[:4].hex()
    except Exception as e:
        print(f"Error processing binary file {file_path}: {e}")
        return None

    # Generate plain text report
    report = [f"{key}: {value}" for key, value in file_info.items() if value]

    return "\n".join(report)

def extract_metadata(file_path):
    """
    Extracts comprehensive metadata from a file of any type.
    
    Args:
        file_path (str): Path to the file to extract metadata from
        
    Returns:
        dict: A dictionary containing all available metadata for the file
    """
    file_path = clean_path(file_path)
    if not file_path:
        return {"error": "File not found or inaccessible"}
    
    # Initialize metadata container
    metadata = {
        "file": {
            "name": os.path.basename(file_path),
            "directory": os.path.dirname(os.path.abspath(file_path)),
            "absolute_path": os.path.abspath(file_path),
            "size_bytes": 0,
            "mime_type": "",
            "extension": os.path.splitext(file_path)[1].lower(),
            "created": "",
            "modified": "",
            "accessed": "",
            "permissions": "",
            "owner": "",
            "hashes": {},
        },
        "system": {
            "hostname": socket.gethostname(),
            "platform": platform.system(),
            "platform_version": platform.version(),
            "processor": platform.processor(),
            "extraction_time": datetime.now().isoformat(),
        },
        "exif": {},
        "extended": {}
    }
    
    # Get basic file stats
    try:
        file_stats = os.stat(file_path)
        metadata["file"]["size_bytes"] = file_stats.st_size
        metadata["file"]["created"] = datetime.fromtimestamp(file_stats.st_ctime).isoformat()
        metadata["file"]["modified"] = datetime.fromtimestamp(file_stats.st_mtime).isoformat()
        metadata["file"]["accessed"] = datetime.fromtimestamp(file_stats.st_atime).isoformat()
        metadata["file"]["permissions"] = oct(file_stats.st_mode & 0o777)
        try:
            import pwd
            metadata["file"]["owner"] = pwd.getpwuid(file_stats.st_uid).pw_name
        except (ImportError, KeyError):
            # pwd module not available on Windows or owner not found
            metadata["file"]["owner"] = str(file_stats.st_uid)
    except Exception as e:
        metadata["file"]["error"] = f"Failed to get file stats: {str(e)}"
    
    # Get MIME type
    try:
        metadata["file"]["mime_type"] = magic.from_file(file_path, mime=True)
    except Exception as e:
        metadata["file"]["error"] = f"Failed to determine MIME type: {str(e)}"
    
    # Calculate file hashes
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            metadata["file"]["hashes"]["md5"] = hashlib.md5(content).hexdigest()
            metadata["file"]["hashes"]["sha1"] = hashlib.sha1(content).hexdigest()
            metadata["file"]["hashes"]["sha256"] = hashlib.sha256(content).hexdigest()
            
            # Calculate entropy
            if content:
                counter = Counter(content)
                length = len(content)
                entropy = -sum((count / length) * math.log2(count / length) for count in counter.values())
                metadata["file"]["entropy"] = entropy
    except Exception as e:
        metadata["file"]["error"] = f"Failed to calculate hashes: {str(e)}"
    
    # Get EXIF data if available
    exif_data = extract_exif(file_path)
    if exif_data:
        metadata["exif"] = exif_data
    
    # Extract extended metadata based on file type
    mime_type = metadata["file"]["mime_type"]
    
    # PDF metadata
    if mime_type == "application/pdf":
        try:
            doc = fitz.open(file_path)
            metadata["extended"]["pdf"] = doc.metadata
            metadata["extended"]["page_count"] = len(doc)
            
            # Get form fields if any
            form_fields = []
            for page in doc:
                fields = page.widgets()
                if fields:
                    for field in fields:
                        form_fields.append({
                            "name": field.field_name,
                            "type": field.field_type,
                            "value": field.field_value
                        })
            
            if form_fields:
                metadata["extended"]["form_fields"] = form_fields
                
            doc.close()
        except Exception as e:
            metadata["extended"]["error"] = f"Error extracting PDF metadata: {str(e)}"
    
    # Word document metadata
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = Document(file_path)
            core_properties = doc.core_properties
            
            # Extract core properties
            metadata["extended"]["word"] = {
                "author": core_properties.author,
                "category": core_properties.category,
                "comments": core_properties.comments,
                "content_status": core_properties.content_status,
                "created": core_properties.created.isoformat() if core_properties.created else None,
                "identifier": core_properties.identifier,
                "keywords": core_properties.keywords,
                "language": core_properties.language,
                "last_modified_by": core_properties.last_modified_by,
                "last_printed": core_properties.last_printed.isoformat() if core_properties.last_printed else None,
                "modified": core_properties.modified.isoformat() if core_properties.modified else None,
                "revision": core_properties.revision,
                "subject": core_properties.subject,
                "title": core_properties.title,
                "version": core_properties.version
            }
            
            # Count paragraphs, tables, and images
            metadata["extended"]["word"]["paragraph_count"] = len(doc.paragraphs)
            metadata["extended"]["word"]["table_count"] = len(doc.tables)
            
            # Count images
            image_count = 0
            for rel in doc.part.rels:
                if "image" in doc.part.rels[rel].target_ref:
                    image_count += 1
            metadata["extended"]["word"]["image_count"] = image_count
            
        except Exception as e:
            metadata["extended"]["error"] = f"Error extracting Word metadata: {str(e)}"
    
    # Old Word DOC format
    elif mime_type == "application/msword":
        try:
            # Use the existing function to extract legacy DOC metadata
            with open(file_path, 'rb') as f:
                binary_data = f.read()
                
            def extract_printable_strings(binary_data):
                pattern = re.compile(b'[' + re.escape(bytes(string.printable, 'ascii')) + b']{4,}')
                matches = pattern.findall(binary_data)
                decoded = [m.decode(errors='ignore').strip() for m in matches]
                return list(dict.fromkeys(decoded))
            
            strings = extract_printable_strings(binary_data)
            
            # Extract metadata from strings
            doc_metadata = {}
            for line in strings:
                if "Microsoft Word" in line:
                    doc_metadata["Software"] = line
                elif "Word.Document" in line:
                    doc_metadata["Format"] = line
                elif "Title" in line and len(line) < 100:
                    doc_metadata["Title"] = line
                elif "Author" in line and len(line) < 100:
                    doc_metadata["Author"] = line
                elif "Company" in line and len(line) < 100:
                    doc_metadata["Company"] = line
                elif "Template" in line and len(line) < 100:
                    doc_metadata["Template"] = line
            
            metadata["extended"]["msword"] = doc_metadata
        except Exception as e:
            metadata["extended"]["error"] = f"Error extracting legacy Word metadata: {str(e)}"
    
    # Excel metadata
    elif mime_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
        try:
            import openpyxl
            if mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                
                excel_metadata = {
                    "sheet_names": workbook.sheetnames,
                    "sheet_count": len(workbook.sheetnames),
                    "properties": {}
                }
                
                # Extract document properties if available
                if hasattr(workbook, 'properties'):
                    props = workbook.properties
                    excel_metadata["properties"] = {
                        "author": props.creator,
                        "title": props.title,
                        "subject": props.subject,
                        "keywords": props.keywords,
                        "category": props.category,
                        "created": props.created.isoformat() if props.created else None,
                        "modified": props.modified.isoformat() if props.modified else None,
                        "lastModifiedBy": props.lastModifiedBy
                    }
                
                # Sheet statistics
                sheet_stats = []
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    stats = {
                        "name": sheet_name,
                        "max_row": sheet.max_row,
                        "max_column": sheet.max_column
                    }
                    sheet_stats.append(stats)
                
                excel_metadata["sheets"] = sheet_stats
                metadata["extended"]["excel"] = excel_metadata
        except Exception as e:
            metadata["extended"]["error"] = f"Error extracting Excel metadata: {str(e)}"
    
    # Image metadata (beyond EXIF)
    elif mime_type.startswith("image/"):
        try:
            img = Image.open(file_path)
            image_metadata = {
                "format": img.format,
                "mode": img.mode,
                "width": img.width,
                "height": img.height,
                "palette": bool(img.palette),
                "info": {k: str(v) for k, v in img.info.items()}
            }
            metadata["extended"]["image"] = image_metadata
        except Exception as e:
            metadata["extended"]["error"] = f"Error extracting image metadata: {str(e)}"
    
    # Audio metadata
    elif mime_type.startswith("audio/"):
        try:
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            audio = AudioSegment.from_file(file_path, format=ext)
            audio_metadata = {
                "channels": audio.channels,
                "sample_width_bytes": audio.sample_width,
                "frame_rate_hz": audio.frame_rate,
                "duration_seconds": len(audio) / 1000.0,
                "frame_count": len(audio.get_array_of_samples()) // audio.channels,
                "max_amplitude": max(audio.get_array_of_samples()) if audio.get_array_of_samples() else 0,
                "bitrate": audio.frame_rate * audio.sample_width * 8 * audio.channels
            }
            metadata["extended"]["audio"] = audio_metadata
        except Exception as e:
            metadata["extended"]["error"] = f"Error extracting audio metadata: {str(e)}"
    
    # Clean up metadata by removing None values
    def clean_dict(d):
        if not isinstance(d, dict):
            return d
        return {k: clean_dict(v) for k, v in d.items() if v is not None}
    
    metadata = clean_dict(metadata)
    
    return metadata


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Extract text or metadata from a file")
    parser.add_argument("file", type=str, help="Path to the input file")
    parser.add_argument("--metadata", action="store_true", help="Extract metadata instead of text")

    args = parser.parse_args()
    file_path = args.file

    if args.metadata:
        result = extract_metadata(file_path)
        if result:
            print("\n--- Extracted Metadata ---\n")
            print(json.dumps(result, indent=2))
        else:
            print("No metadata could be extracted or an error occurred.")
    else:
        result = extract_text(file_path)
        if result:
            print("\n--- Extracted Text ---\n")
            print(result)
        else:
            print("No text could be extracted or an error occurred.")

if __name__ == "__main__":
    main()
